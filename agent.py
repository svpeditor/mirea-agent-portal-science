"""science-agent — поиск РЕАЛЬНЫХ научных статей по теме.

Принцип: агент НИКОГДА не выдумывает публикации. LLM используется только
чтобы (1) сгенерировать поисковые запросы, (2) оценить релевантность реально
найденных статей и перевести/написать аннотации по их реальным abstract'ам.
Сами статьи всегда берутся из реальных индексов через allowlist-прокси
portal-api:

- CyberLeninka     (/api/sandbox/cyberleninka)      — русскоязычные журналы
- arXiv            (/api/sandbox/arxiv)             — препринты STEM
- Crossref         (/api/sandbox/crossref)          — DOI-журналы
- Semantic Scholar (/api/sandbox/semantic-scholar)  — abstracts + цитируемость
- OpenAlex         (/api/sandbox/openalex)          — широкое покрытие + OA

Результаты всех источников по всем запросам объединяются и дедуплицируются.
Если реальный поиск ничего не дал — агент честно завершается с ошибкой, а НЕ
подменяет выдачу галлюцинацией (никакого «режима знаний LLM»).

PDF открытого доступа (arXiv, CyberLeninka, OA по DOI через Unpaywall)
скачивается в output/pdfs/. Только легально открытые версии.
"""
from __future__ import annotations

import json
import os
import re
import time
from dataclasses import dataclass, field

import httpx
from docx import Document

from portal_sdk import Agent

MAX_PDF_DOWNLOADS = 10  # сколько верхних статей качать файлом
# Совокупный бюджет PDF: заведомо ниже portal max_job_output_bytes (1 GiB),
# чтобы report.docx/sources.bib всегда сохранились, даже если PDF большие.
MAX_PDF_TOTAL_BYTES = 200 * 1024 * 1024
MAX_QUERIES = 5            # сколько поисковых запросов генерим
RESULTS_PER_QUERY = 10     # сколько результатов берём из каждого источника на запрос
MAX_FOR_ANALYSIS = 40      # сколько верхних статей отправляем на LLM-оценку


@dataclass
class Paper:
    title: str
    authors: list[str]
    year: int | None
    venue: str
    arxiv_id: str | None
    url: str | None
    annotation: str
    score: float = 0.0
    score_explanation: str = ""
    pdf_url: str | None = None
    doi: str | None = None
    cl_slug: str | None = None  # slug статьи CyberLeninka (для скачивания PDF)
    citation_count: int = 0
    provenance: list[str] = field(default_factory=list)  # ['arXiv','CyberLeninka',...]
    pdf_filename: str | None = None  # имя файла в output/pdfs/ если скачали

    @property
    def bibkey(self) -> str:
        parts = (self.authors[0] if self.authors else "anon").split()
        surname = max(parts, key=len) if parts else "anon"
        surname = re.sub(r"[^a-z0-9]", "", surname.lower()) or "anon"
        # ASCII-срез заголовка; для кириллических заголовков он пуст, поэтому
        # берём запасной идентификатор (arxiv/doi/slug), чтобы ключи не
        # схлопывались в одинаковый «фамилияГод_».
        slug = re.sub(r"[^A-Za-z0-9]", "", self.title)[:20].lower()
        if not slug:
            alt = self.arxiv_id or self.doi or self.cl_slug or ""
            slug = re.sub(r"[^a-z0-9]", "", alt.lower())[:16]
        return f"{surname}{self.year or 'nd'}_{slug}"

    @property
    def best_link(self) -> str | None:
        aid = _norm_arxiv_id(self.arxiv_id)
        if aid:
            return f"https://arxiv.org/abs/{aid}"
        if self.doi:
            return f"https://doi.org/{self.doi}"
        return self.url


def _llm_call(messages: list[dict], model: str, api_key: str, base_url: str, *, max_tokens: int = 8000) -> str:
    payload = {"model": model, "messages": messages, "temperature": 0.3, "max_tokens": max_tokens}
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://github.com/svpeditor/mirea-agent-portal-science",
        "X-Title": "mirea-science-agent",
    }
    with httpx.Client(timeout=300) as client:
        r = client.post(f"{base_url.rstrip('/')}/chat/completions", json=payload, headers=headers)
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"]


def _strip_reasoning(s: str) -> str:
    s = re.sub(r"<think>.*?</think>", "", s, flags=re.DOTALL).strip()
    m = re.search(r"```(?:json)?\s*(.*?)\s*```", s, re.DOTALL)
    if m:
        s = m.group(1)
    return s.strip()


def _parse_json(s: str):
    s = _strip_reasoning(s)
    m = re.search(r"[\[\{].*[\]\}]", s, re.DOTALL)
    if m:
        s = m.group(0)
    return json.loads(s)


def _norm_title(t: str) -> str:
    return re.sub(r"[^a-zа-я0-9]+", "", (t or "").lower())


def _has_cyrillic(s: str) -> bool:
    return bool(re.search(r"[а-яА-Я]", s or ""))


def _strip_jats(s: str) -> str:
    """Crossref abstract приходит JATS-XML — выкидываем теги."""
    return re.sub(r"<[^>]+>", " ", s or "").strip()


def _sandbox_root(base_url: str) -> str:
    base = base_url.rstrip("/")
    if base.endswith("/llm/v1"):
        base = base[: -len("/llm/v1")]
    return base


def _norm_arxiv_id(s: str | None) -> str | None:
    """Единственная точка нормализации arxiv_id: срезаем префикс `arxiv:`
    и пробелы. Используется везде (ingest/ссылки/скачивание)."""
    if not s:
        return None
    out = re.sub(r"(?i)^arxiv:", "", s).strip()
    return out or None


def _surname(authors: list[str]) -> str:
    parts = (authors[0] if authors else "").split()
    return max(parts, key=len).lower() if parts else ""


def _arxiv_pdf_url(arxiv_id: str | None) -> str | None:
    aid = _norm_arxiv_id(arxiv_id)
    return f"https://arxiv.org/pdf/{aid}.pdf" if aid else None


# --- Генерация поисковых запросов через LLM ---

def _gen_queries(topic: str, description: str, model: str, api_key: str, base_url: str) -> list[str]:
    """LLM генерит короткие запросы RU+EN. На сбое — fallback на саму тему."""
    system = (
        "Ты — эксперт по научному информационному поиску. Сгенерируй короткие "
        "(2-6 слов) поисковые запросы для научных баз. Часть на РУССКОМ (для "
        "CyberLeninka), часть на АНГЛИЙСКОМ (для международных баз). И широкие, "
        "и узкие. Ответь строго JSON без markdown: {\"queries\": [\"запрос\", ...]}"
    )
    desc = f"Описание: {description}" if description else "Описание не дано."
    user = f"Тема: {topic}\n{desc}\nСгенерируй {MAX_QUERIES} запросов."
    try:
        data = _parse_json(_llm_call(
            [{"role": "system", "content": system}, {"role": "user", "content": user}],
            model=model, api_key=api_key, base_url=base_url, max_tokens=2000,
        ))
        raw = data.get("queries", []) if isinstance(data, dict) else data
        out: list[str] = []
        for q in raw:
            q = str(q).strip()
            if q and q.lower() not in {x.lower() for x in out}:
                out.append(q)
        out = out[:MAX_QUERIES]
        if out:
            if topic not in out:
                out.insert(0, topic)
            return out[:MAX_QUERIES]
    except Exception:  # noqa: BLE001
        pass
    return [topic]


# --- Реальные источники через sandbox-прокси ---

def _get_json(client: httpx.Client, url: str, params: dict, api_key: str) -> dict:
    r = client.get(url, params=params, headers={"Authorization": f"Bearer {api_key}"})
    r.raise_for_status()
    return r.json()


def _search_arxiv(query: str, n: int, api_key: str, root: str) -> list[Paper]:
    out: list[Paper] = []
    with httpx.Client(timeout=60) as c:
        data = _get_json(c, f"{root}/api/sandbox/arxiv",
                         {"search_query": query, "max_results": n}, api_key)
    for row in data.get("papers", []):
        aid = _norm_arxiv_id(row.get("arxiv_id"))
        out.append(Paper(
            title=row.get("title", ""), authors=row.get("authors", []),
            year=row.get("year"), venue="arXiv", arxiv_id=aid,
            url=row.get("url"), annotation=_strip_jats(row.get("abstract", "")),
            pdf_url=_arxiv_pdf_url(aid), provenance=["arXiv"],
        ))
    return out


def _search_crossref(query: str, n: int, api_key: str, root: str) -> list[Paper]:
    out: list[Paper] = []
    with httpx.Client(timeout=60) as c:
        data = _get_json(c, f"{root}/api/sandbox/crossref",
                         {"query": query, "rows": n}, api_key)
    for row in data.get("works", []):
        out.append(Paper(
            title=row.get("title", ""), authors=row.get("authors", []),
            year=row.get("year"), venue=row.get("venue", ""), arxiv_id=None,
            url=row.get("url"), annotation=_strip_jats(row.get("abstract", "")),
            doi=row.get("doi"), citation_count=int(row.get("citation_count") or 0),
            provenance=["Crossref"],
        ))
    return out


def _search_s2(query: str, n: int, api_key: str, root: str) -> list[Paper]:
    out: list[Paper] = []
    with httpx.Client(timeout=60) as c:
        data = _get_json(c, f"{root}/api/sandbox/semantic-scholar",
                         {"query": query, "limit": n}, api_key)
    for row in data.get("papers", []):
        aid = _norm_arxiv_id(row.get("arxiv_id"))
        out.append(Paper(
            title=row.get("title", ""), authors=row.get("authors", []),
            year=row.get("year"), venue=row.get("venue", ""), arxiv_id=aid,
            url=row.get("url"), annotation=_strip_jats(row.get("abstract", "")),
            doi=row.get("doi"), citation_count=int(row.get("citation_count") or 0),
            pdf_url=_arxiv_pdf_url(aid), provenance=["Semantic Scholar"],
        ))
    return out


def _search_cyberleninka(query: str, n: int, api_key: str, root: str) -> list[Paper]:
    out: list[Paper] = []
    with httpx.Client(timeout=60) as c:
        data = _get_json(c, f"{root}/api/sandbox/cyberleninka",
                         {"query": query, "size": n}, api_key)
    for row in data.get("articles", []):
        out.append(Paper(
            title=row.get("title", ""), authors=row.get("authors", []),
            year=row.get("year"), venue=row.get("journal", ""), arxiv_id=None,
            url=row.get("url"), annotation=row.get("abstract", ""),
            cl_slug=row.get("slug") or None, pdf_url=row.get("pdf_url") or None,
            provenance=["CyberLeninka"],
        ))
    return out


def _search_openalex(query: str, n: int, api_key: str, root: str) -> list[Paper]:
    out: list[Paper] = []
    with httpx.Client(timeout=60) as c:
        data = _get_json(c, f"{root}/api/sandbox/openalex",
                         {"query": query, "per_page": n}, api_key)
    for row in data.get("works", []):
        out.append(Paper(
            title=row.get("title", ""), authors=row.get("authors", []),
            year=row.get("year"), venue=row.get("venue", ""), arxiv_id=None,
            url=row.get("url"), annotation=row.get("abstract", ""),
            doi=row.get("doi"), citation_count=int(row.get("citation_count") or 0),
            pdf_url=row.get("oa_url") if row.get("is_oa") else None,
            provenance=["OpenAlex"],
        ))
    return out


# (source_name, fn, only_cyrillic_queries)
_SOURCES = (
    ("CyberLeninka", _search_cyberleninka, True),
    ("arXiv", _search_arxiv, False),
    ("Crossref", _search_crossref, False),
    ("Semantic Scholar", _search_s2, False),
    ("OpenAlex", _search_openalex, False),
)


def _merge(groups: list[list[Paper]]) -> list[Paper]:
    """Дедуп. Сильное тождество — DOI: тогда сливаем И идентификаторы.
    Слабое — нормализованное название + год + фамилия первого автора:
    тогда сливаем ТОЛЬКО provenance и описательные поля, идентификаторы
    (arxiv_id/doi/url/pdf_url/cl_slug) НЕ заимствуем — иначе можно подвесить
    чужой PDF к статье (misattribution)."""
    by_key: dict[str, Paper] = {}
    for group in groups:
        for p in group:
            if not p.title.strip():
                continue
            doi = (p.doi or "").strip().lower()
            if doi:
                key, strong = f"doi:{doi}", True
            else:
                nt = _norm_title(p.title)
                if not nt:
                    continue
                key = f"t:{nt}|{p.year or ''}|{_surname(p.authors)}"
                strong = False
            if key not in by_key:
                by_key[key] = p
                continue
            ex = by_key[key]
            for src in p.provenance:
                if src not in ex.provenance:
                    ex.provenance.append(src)
            ex.year = ex.year or p.year
            ex.venue = ex.venue or p.venue
            ex.citation_count = max(ex.citation_count, p.citation_count)
            if len(p.annotation) > len(ex.annotation):
                ex.annotation = p.annotation
            if len(p.authors) > len(ex.authors):
                ex.authors = p.authors
            if strong:  # только при совпадении по DOI безопасно сливать ID
                ex.arxiv_id = ex.arxiv_id or p.arxiv_id
                ex.doi = ex.doi or p.doi
                ex.url = ex.url or p.url
                ex.pdf_url = ex.pdf_url or p.pdf_url
                ex.cl_slug = ex.cl_slug or p.cl_slug
    return list(by_key.values())


def _prelim_rank(papers: list[Paper]) -> list[Paper]:
    """Эвристика до LLM-анализа: свежесть + цитируемость + наличие аннотации."""
    def s(p: Paper) -> float:
        v = 0.0
        if p.citation_count:
            v += min(p.citation_count / 100, 10)
        if p.annotation:
            v += 2
        if p.year and p.year >= 2020:
            v += 3
        elif p.year and p.year >= 2015:
            v += 1
        return v
    return sorted(papers, key=s, reverse=True)


def _llm_rank_and_annotate(topic, papers, language, model, api_key, base_url) -> list[Paper]:
    """LLM оценивает РЕАЛЬНЫЕ статьи и пишет аннотации по их abstract'ам.
    Не добавляет и не удаляет статьи — только score/annotation/explanation."""
    if not papers:
        return papers
    ann_lang = "русском" if language == "ru" else "английском"
    items = [
        {"i": i, "title": p.title, "year": p.year, "authors": p.authors[:3],
         "abstract": (p.annotation or "")[:600]}
        for i, p in enumerate(papers)
    ]
    system = (
        "Ты — научный библиограф. Дан список РЕАЛЬНЫХ статей. Для каждой: "
        f"(1) score 0..1 релевантности теме, (2) аннотация на {ann_lang} ~5 "
        "предложений строго по приведённому abstract (не выдумывай фактов), "
        "(3) объяснение балла 1-2 предложения. Строго JSON-массив."
    )
    user = (
        f"Тема: {topic}\n\nСтатьи: {json.dumps(items, ensure_ascii=False)}\n\n"
        'Верни [{"i","score","annotation","score_explanation"}]. Без markdown.'
    )
    try:
        parsed = _parse_json(_llm_call(
            [{"role": "system", "content": system}, {"role": "user", "content": user}],
            model=model, api_key=api_key, base_url=base_url, max_tokens=6000,
        ))
    except Exception:  # noqa: BLE001
        return papers  # LLM сорвалась — отдаём реальные статьи как есть
    if not isinstance(parsed, list):
        return papers
    seen: set[int] = set()
    for row in parsed:
        if not isinstance(row, dict):
            continue
        idx = row.get("i")
        if not isinstance(idx, int) or idx < 0 or idx >= len(papers) or idx in seen:
            continue
        seen.add(idx)
        p = papers[idx]
        p.score = float(row.get("score", 0.0) or 0.0)
        new_ann = str(row.get("annotation", "")).strip()
        if new_ann:
            p.annotation = new_ann
        p.score_explanation = str(row.get("score_explanation", "")).strip()
    return papers


def _download_pdfs(agent: Agent, papers: list[Paper], api_key: str, root: str) -> None:
    """Качаем PDF открытого доступа в output/pdfs/ через sandbox-прокси.

    Приоритет источника: arXiv -> CyberLeninka -> OA по DOI (Unpaywall).
    Только легально открытые версии; на любой проблеме оставляем ссылку.
    """
    pdf_dir = agent.output_dir / "pdfs"
    done = 0
    total = 0
    for p in papers:
        if done >= MAX_PDF_DOWNLOADS or total >= MAX_PDF_TOTAL_BYTES:
            break
        aid = _norm_arxiv_id(p.arxiv_id)
        if aid:
            endpoint, params, fname, tag = (
                "arxiv-pdf", {"arxiv_id": aid},
                f"{re.sub(r'[^A-Za-z0-9._-]', '_', aid)}.pdf", aid,
            )
        elif p.cl_slug:
            endpoint, params, fname, tag = (
                "cyberleninka-pdf", {"slug": p.cl_slug},
                f"cl_{re.sub(r'[^A-Za-z0-9._-]', '_', p.cl_slug)}.pdf", p.cl_slug,
            )
        elif p.doi:
            endpoint, params, fname, tag = (
                "oa-pdf", {"doi": p.doi},
                f"{re.sub(r'[^A-Za-z0-9._-]', '_', p.doi)}.pdf", p.doi,
            )
        else:
            continue
        try:
            with httpx.Client(timeout=90) as c:
                r = c.get(f"{root}/api/sandbox/{endpoint}", params=params,
                          headers={"Authorization": f"Bearer {api_key}"})
            if r.status_code != 200:
                agent.log("warn", f"PDF {tag}: прокси {r.status_code}, оставляю ссылку")
                continue
            if total + len(r.content) > MAX_PDF_TOTAL_BYTES:
                agent.log("warn", f"PDF {tag}: превышен общий бюджет, оставляю ссылку")
                break
            pdf_dir.mkdir(parents=True, exist_ok=True)
            (pdf_dir / fname).write_bytes(r.content)
            p.pdf_filename = f"pdfs/{fname}"
            done += 1
            total += len(r.content)
        except Exception as e:  # noqa: BLE001
            agent.log("warn", f"PDF {tag}: не скачал ({e}), оставляю ссылку")
    if done:
        agent.log("info", f"Скачано PDF-файлов: {done} ({total // 1024} КБ)")


def _build_bibtex(papers: list[Paper]) -> str:
    out: list[str] = []
    used: dict[str, int] = {}
    for p in papers:
        # Гарантируем уникальность ключа (иначе LaTeX падает на дублях).
        base = p.bibkey
        used[base] = used.get(base, 0) + 1
        key = base if used[base] == 1 else f"{base}_{used[base]}"
        authors = " and ".join(p.authors) if p.authors else "Unknown"
        fields = [
            f"  title = {{{p.title}}}",
            f"  author = {{{authors}}}",
            f"  year = {{{p.year or 'n.d.'}}}",
        ]
        if p.venue:
            fields.append(f"  journal = {{{p.venue}}}")
        if p.arxiv_id:
            fields.append(f"  eprint = {{arxiv:{p.arxiv_id}}}")
        if p.doi:
            fields.append(f"  doi = {{{p.doi}}}")
        if p.url:
            fields.append(f"  url = {{{p.url}}}")
        out.append("@article{" + key + ",\n" + ",\n".join(fields) + "\n}\n")
    return "\n".join(out)


def _build_report(topic: str, papers: list[Paper], model: str, queries: list[str]) -> Document:
    doc = Document()
    doc.add_heading("Поиск научных статей — отчёт", level=0)
    doc.add_paragraph(f"Тема: {topic}")
    doc.add_paragraph(f"Всего статей: {len(papers)}")
    doc.add_paragraph(f"Модель: {model}")

    srcs = sorted({s for p in papers for s in p.provenance})
    doc.add_paragraph(
        "Все статьи ниже — реальные записи из баз: "
        + (", ".join(srcs) if srcs else "—")
        + ". Ссылки кликабельны; PDF открытого доступа приложены файлами."
    )

    if queries:
        doc.add_heading("Поисковые запросы", level=1)
        for q in queries:
            doc.add_paragraph(q, style="List Bullet")

    doc.add_heading("Ранжированный список", level=1)
    for i, p in enumerate(papers, start=1):
        doc.add_heading(f"{i}. {p.title}", level=2)
        if p.authors:
            doc.add_paragraph(
                f"Авторы: {', '.join(p.authors[:6])}"
                f"{' и др.' if len(p.authors) > 6 else ''}"
            )
        meta = []
        if p.year:
            meta.append(str(p.year))
        if p.venue:
            meta.append(p.venue)
        if p.arxiv_id:
            meta.append(f"arXiv:{p.arxiv_id}")
        if p.doi:
            meta.append(f"DOI:{p.doi}")
        if p.citation_count:
            meta.append(f"цитирований: {p.citation_count}")
        if p.score:
            meta.append(f"релевантность {p.score:.2f}")
        if p.provenance:
            meta.append("источник: " + ", ".join(p.provenance))
        if meta:
            doc.add_paragraph("  •  ".join(meta))
        if p.annotation:
            doc.add_paragraph(f"Аннотация: {p.annotation}")
        if p.score_explanation:
            doc.add_paragraph(f"Почему такой балл: {p.score_explanation}")
        link = p.best_link
        if link:
            doc.add_paragraph(f"Ссылка: {link}")
        if p.pdf_filename:
            doc.add_paragraph(f"Файл статьи: {p.pdf_filename}")
        elif p.pdf_url:
            doc.add_paragraph(f"PDF: {p.pdf_url}")
    return doc


def main() -> None:
    agent = Agent()
    params = agent.params
    topic: str = (params.get("topic") or "").strip()
    description: str = (params.get("description") or "").strip()
    max_papers: int = max(5, min(30, int(params.get("max_papers", 15))))
    language: str = params.get("language", "ru")
    sort_by: str = (params.get("sort_by") or "relevance").lower()
    download_pdf: bool = bool(params.get("download_pdf", True))

    api_key = os.environ.get("OPENROUTER_API_KEY", "").strip()
    base_url = os.environ.get("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1").strip()
    model = (params.get("model") or os.environ.get("LLM_MODEL") or "deepseek/deepseek-r1").strip()
    root = _sandbox_root(base_url)

    if not api_key:
        agent.failed("OPENROUTER_API_KEY не передан.")
        return
    if not topic:
        agent.failed("Не указана тема.")
        return

    agent.log("info", f"science: topic={topic[:80]!r}, max={max_papers}, "
                       f"lang={language}, model={model}, sort={sort_by}")

    agent.progress(0.1, f"{model}: генерация поисковых запросов")
    queries = _gen_queries(topic, description, model, api_key, base_url)
    agent.log("info", f"Запросы ({len(queries)}): {queries}")

    groups: list[list[Paper]] = []
    errors: list[str] = []
    n_src = len(_SOURCES)
    for si, (name, fn, ru_only) in enumerate(_SOURCES):
        # CyberLeninka — только по русским запросам (если есть); остальным — все.
        src_queries = [q for q in queries if _has_cyrillic(q)] if ru_only else queries
        if ru_only and not src_queries:
            src_queries = queries  # русских нет — пробуем как есть
        found = 0
        for q in src_queries:
            agent.progress(0.15 + 0.45 * si / n_src, f"Поиск: {name} «{q[:40]}»")
            try:
                res = fn(q, RESULTS_PER_QUERY, api_key, root)
                groups.append(res)
                found += len(res)
            except Exception as e:  # noqa: BLE001
                errors.append(f"{name}: {e}")
                agent.log("warn", f"{name} «{q[:40]}»: {e}")
            time.sleep(0.4)  # бережём rate-limit апстримов
        agent.log("info", f"{name}: найдено {found}")

    papers = _merge(groups)
    if not papers:
        if not any(g for g in groups) and errors:
            agent.failed(
                "Источники поиска недоступны (ни один не ответил) — это сбой "
                "инфраструктуры, а не отсутствие статей. Повторите запуск позже. "
                "Детали: " + "; ".join(errors[:5])
            )
        else:
            detail = "; ".join(errors[:5]) if errors else "все источники вернули 0 результатов"
            agent.failed(
                "По теме не найдено реальных публикаций в CyberLeninka / arXiv / "
                "Crossref / Semantic Scholar / OpenAlex. Уточните формулировку "
                f"(добавьте ключевые слова). Детали: {detail}"
            )
        return

    candidates = _prelim_rank(papers)[:MAX_FOR_ANALYSIS]
    agent.progress(0.65, f"{model}: оценка релевантности и аннотации ({len(candidates)})")
    candidates = _llm_rank_and_annotate(topic, candidates, language, model, api_key, base_url)

    if sort_by == "popularity":
        candidates.sort(key=lambda p: (-p.citation_count, -p.score))
    else:
        candidates.sort(key=lambda p: (-p.score, -p.citation_count))
    papers = candidates[:max_papers]

    if download_pdf:
        agent.progress(0.85, "Скачиваю PDF открытого доступа")
        _download_pdfs(agent, papers, api_key, root)

    for p in papers:
        agent.item_done(p.arxiv_id or p.doi or p.title[:40], summary=p.title,
                        data={"year": p.year, "venue": p.venue, "score": p.score})

    agent.progress(0.95, "Формирую report.docx и sources.bib")
    out_dir = agent.output_dir
    _build_report(topic, papers, model, queries).save(out_dir / "report.docx")
    (out_dir / "sources.bib").write_text(_build_bibtex(papers), encoding="utf-8")

    agent.progress(1.0, "Готово")
    agent.result(artifacts=[
        {"id": "report", "path": "report.docx"},
        {"id": "bibtex", "path": "sources.bib"},
    ])


if __name__ == "__main__":
    main()
